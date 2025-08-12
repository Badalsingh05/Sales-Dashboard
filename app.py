import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt
from fpdf import FPDF
from docx import Document
from io import BytesIO

# Load and clean data
df = pd.read_csv("sales_data.csv")
df = df[df["Date"] != "Date"]
df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
df["Sales"] = pd.to_numeric(df["Sales"], errors="coerce")
df["Quantity"] = pd.to_numeric(df["Quantity"], errors="coerce")
df.dropna(subset=["Date", "Sales", "Quantity"], inplace=True)

# Page config
st.set_page_config(page_title="Sales Dashboard", layout="wide")

# Sidebar filters
st.sidebar.header("Filter Data")
regions = st.sidebar.multiselect("Region", df["Region"].unique(), default=df["Region"].unique())
products = st.sidebar.multiselect("Product", df["Product"].unique(), default=df["Product"].unique())
date_range = st.sidebar.date_input("Date Range", [df["Date"].min(), df["Date"].max()])

# Filter data
filtered_df = df[
    (df["Region"].isin(regions)) &
    (df["Product"].isin(products)) &
    (df["Date"] >= pd.to_datetime(date_range[0])) &
    (df["Date"] <= pd.to_datetime(date_range[1]))
]

# Title
st.title("Sales Dashboard")
st.markdown("Responsive dashboard for tracking sales across regions and products.")

# KPI Cards
k1, k2, k3 = st.columns(3)
k1.metric("Total Sales", f"${filtered_df['Sales'].sum():,.2f}")
k2.metric("Total Quantity", int(filtered_df["Quantity"].sum()))
k3.metric("Transactions", len(filtered_df))

st.markdown("---")

# Charts
c1, c2 = st.columns([2, 1])

with c1:
    st.subheader("Sales Over Time")
    sales_over_time = filtered_df.groupby("Date")["Sales"].sum().reset_index()
    st.line_chart(sales_over_time.set_index("Date"))

with c2:
    st.subheader("Sales by Region")
    sales_by_region = filtered_df.groupby("Region")["Sales"].sum().sort_values()
    st.bar_chart(sales_by_region)

st.markdown("### Sales by Product")
sales_by_product = filtered_df.groupby("Product")["Sales"].sum().sort_values()
st.bar_chart(sales_by_product)

# Raw Data
with st.expander("Show Raw Data Table"):
    st.dataframe(filtered_df, use_container_width=True)


st.markdown("---")
st.subheader("⬇ Export Filtered Data")

# Convert to CSV
csv = filtered_df.to_csv(index=False).encode("utf-8")
st.download_button(" Download CSV", data=csv, file_name="filtered_data.csv", mime="text/csv")

# Convert to Excel
excel_buffer = BytesIO()
with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
    filtered_df.to_excel(writer, index=False, sheet_name="Sales Data")
st.download_button("Download Excel", data=excel_buffer.getvalue(), file_name="filtered_data.xlsx")

# Convert to PDF
def create_pdf(df):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=10)
    pdf.cell(200, 10, txt="Filtered Sales Report", ln=True, align="C")

    col_width = pdf.w / 5.5
    row_height = 8

    pdf.set_font("Arial", size=8)
    pdf.set_fill_color(220, 220, 220)

    for i, column in enumerate(df.columns):
        pdf.cell(col_width, row_height, txt=column, border=1, fill=True)

    pdf.ln(row_height)

    for index, row in df.iterrows():
        for item in row:
            pdf.cell(col_width, row_height, txt=str(item), border=1)
        pdf.ln(row_height)
    return pdf.output(dest="S").encode("latin1")

pdf_data = create_pdf(filtered_df)
st.download_button("Download PDF", data=pdf_data, file_name="filtered_data.pdf", mime="application/pdf")

# Convert to Word
def create_word(df):
    doc = Document()
    doc.add_heading("Filtered Sales Report", 0)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)
    word_file = BytesIO()
    doc.save(word_file)
    return word_file.getvalue()

word_data = create_word(filtered_df)
st.download_button("Download Word", data=word_data, file_name="filtered_data.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# Footer
st.markdown("---")
st.caption("Fully responsive. Built with ❤️ Mr.Badal")
